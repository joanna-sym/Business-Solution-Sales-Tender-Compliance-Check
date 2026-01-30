"""
Filename: step00_generate_mock_data.py
Description: [Step 00] Generates mock tender documents (.docx) and product specs (.xlsx) for testing.
             Run this script first to prepare the environment.
Author: Joanna
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
import os

def create_mock_files():
    print("[START] Step 00: Initializing Data Generation...")

    # --- 1. Generate Tender Document (Word) ---
    doc_name = "mock_tender_document.docx"
    doc = Document()
    
    # Title
    doc.add_heading('Medical Device Procurement Project - 2026', 0)
    
    # Mandatory Clause (The Trap)
    doc.add_heading('Section 1: Qualification Requirements', level=1)
    p = doc.add_paragraph()
    run = p.add_run('★ Requirement 1: Valid NMPA Registration Certificate is mandatory. (Rejection Term)')
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0) # Red
    
    # Technical Table
    doc.add_heading('Section 2: Technical Specifications', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Parameter Name'
    hdr_cells[2].text = 'Tender Requirement'
    
    # Data Rows (Containing Traps)
    items = [
        ("1", "Stapler Length", ">= 60mm"),
        ("2", "★ Firing Safety", "Double-lock safety mechanism (Mandatory)"),
        ("3", "Rotation Angle", "360-degree rotatable handle"),
        ("4", "Application", "Must include Pediatric Surgery in certificate")
    ]
    
    for id_val, name, req in items:
        row_cells = table.add_row().cells
        row_cells[0].text = id_val
        row_cells[1].text = name
        row_cells[2].text = req

    doc.save(doc_name)
    print(f"[OK] Generated: {doc_name}")

    # --- 2. Generate Product Specs (Excel) ---
    xls_name = "product_specs.xlsx"
    
    # Mock Product Data (Intentionally fails some checks)
    data = {
        'Parameter Name': [
            'Stapler Length',
            'Firing Safety',
            'Rotation Angle',
            'Application'
        ],
        'Our Product Spec': [
            '58mm',                          # Trap: 58 < 60 (FAIL)
            'Single-lock safety mechanism',  # Trap: Single vs Double (FAIL)
            '360-degree rotatable handle',   # PASS
            'General Surgery only'           # Trap: Missing Pediatric (FAIL)
        ]
    }
    
    df = pd.DataFrame(data)
    df.to_excel(xls_name, index=False)
    print(f"[OK] Generated: {xls_name}")

    print("\n[DONE] All mock data is ready! Now proceed to 'step01_tender_auditor.py'.")

if __name__ == "__main__":
    create_mock_files()